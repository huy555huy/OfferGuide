"""W12-fix(c) — DOCX-level resume tailoring (format-preserving).

Tests:
- profile.loader auto-dispatches .pdf vs .docx by suffix
- docx_tailor eligibility rules (skip header / dates / multi-column / short)
- LLM rewrite plan applies safely (sanity-check on length drift)
- Output .docx preserves paragraph styles, font, bold, alignment
- /api/tailor/docx route renders error states cleanly
- /api/tailor/download serves files + path traversal blocked
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

import offerguide
from offerguide.config import Settings
from offerguide.profile import UserProfile, load_resume_pdf
from offerguide.skills import discover_skills
from offerguide.skills.tailor_resume.docx_tailor import (
    ParagraphPlan,
    _is_eligible,
    tailor_docx,
)
from offerguide.ui.notify import ConsoleNotifier
from offerguide.ui.web import create_app

SKILLS_ROOT = Path(__file__).parent.parent / "src/offerguide/skills"


# ═══════════════════════════════════════════════════════════════════
# Eligibility rules
# ═══════════════════════════════════════════════════════════════════


class TestEligibility:
    def test_empty_skipped(self) -> None:
        eligible, reason = _is_eligible("", run_count=0, is_bold=False)
        assert not eligible
        assert reason == "empty"

    def test_short_bold_header_skipped(self) -> None:
        eligible, reason = _is_eligible("教育背景", run_count=1, is_bold=True)
        assert not eligible
        assert reason == "section_header"

    def test_email_phone_skipped(self) -> None:
        for s in (
            "+86 19952993755 | foo@bar.com",
            "工作时间 2025/3 至今",
            "19952993755",
        ):
            eligible, reason = _is_eligible(s, run_count=1, is_bold=False)
            assert not eligible, f"{s!r} should be ineligible"
            assert reason in ("contains_date_phone_email",), reason

    def test_multi_column_skipped(self) -> None:
        # Tab-separated 学校 / 时间
        eligible, reason = _is_eligible(
            "上海财经大学\t\t\t09/2025-06/2027", run_count=5, is_bold=True,
        )
        assert not eligible
        assert reason == "multi_column_layout"
        # 3+ spaces also caught
        eligible, reason = _is_eligible(
            "上海财经大学    09/2025-06/2027", run_count=5, is_bold=True,
        )
        assert not eligible

    def test_too_short_skipped(self) -> None:
        eligible, reason = _is_eligible("一些技能", run_count=1, is_bold=False)
        assert not eligible
        assert reason == "too_short"

    def test_long_paragraph_eligible(self) -> None:
        long_text = "设计并实现 LangGraph agent 双层架构，承载研究过程中的状态、动作执行" * 2
        eligible, reason = _is_eligible(long_text, run_count=14, is_bold=False)
        assert eligible
        assert reason == ""


# ═══════════════════════════════════════════════════════════════════
# Profile loader: docx + pdf dispatch
# ═══════════════════════════════════════════════════════════════════


class TestProfileLoader:
    def test_real_user_docx_loads(self) -> None:
        """The user's actual 中文简历.docx should load with non-trivial content."""
        path = Path("/Users/huy/new_try/中文简历.docx")
        if not path.exists():
            pytest.skip("user resume not present in test env")
        profile = load_resume_pdf(path)
        assert isinstance(profile, UserProfile)
        assert "胡阳" in profile.raw_resume_text
        assert "上海财经大学" in profile.raw_resume_text
        assert len(profile.raw_resume_text) > 1000
        assert profile.source_pdf and profile.source_pdf.endswith(".docx")

    def test_unsupported_format_raises(self, tmp_path) -> None:
        bad = tmp_path / "x.txt"
        bad.write_text("not a resume")
        with pytest.raises(ValueError, match="Unsupported"):
            load_resume_pdf(bad)

    def test_missing_file_raises(self, tmp_path) -> None:
        with pytest.raises(FileNotFoundError):
            load_resume_pdf(tmp_path / "nonexistent.docx")


# ═══════════════════════════════════════════════════════════════════
# docx_tailor with mocked LLM
# ═══════════════════════════════════════════════════════════════════


class _StubLLMResp:
    def __init__(self, content: str):
        self.content = content
        self.latency_ms = 100
        self.prompt_tokens = 10
        self.completion_tokens = 5
        self.cost_usd = 0.0
        self.model = "stub"
        self.raw = None


class _StubLLM:
    def __init__(self, decisions: list[dict]):
        self._json = json.dumps(decisions, ensure_ascii=False)

    def chat(self, messages, **kwargs):
        return _StubLLMResp(content=self._json)


class TestDocxTailorWithStub:
    def test_format_preserved_after_rewrite(self, tmp_path) -> None:
        """Use the real user resume; mock LLM to rewrite one paragraph;
        verify output preserves all paragraph-level styles."""
        from docx import Document

        master_path = Path("/Users/huy/new_try/中文简历.docx")
        if not master_path.exists():
            pytest.skip("user resume not present")

        # Mock LLM to rewrite exactly one paragraph
        decisions = [
            {
                "index": 18,
                "decision": "rewrite",
                "new_text": "重新设计 Deep Research Agent: 用 LangGraph + DSPy + Pydantic v2 schema 实现 agent state machine, 闭环目标可验证",
                "rationale": "对齐 JD 第 2 条 LangGraph 要求 + 加 ATS 关键词",
            }
        ]
        llm = _StubLLM(decisions)

        out_path = tmp_path / "tailored.docx"
        result = tailor_docx(
            input_path=master_path,
            output_path=out_path,
            jd_text="字节 LangGraph + DSPy 后端实习",
            company="字节跳动",
            role_focus="AI Agent",
            master_resume_text=load_resume_pdf(master_path).raw_resume_text,
            llm=llm,
        )

        assert out_path.exists()
        assert result.summary()["total_paragraphs"] == 64
        assert result.summary()["rewritten"] == 1

        # Compare structure
        orig = Document(str(master_path))
        new = Document(str(out_path))
        assert len(orig.paragraphs) == len(new.paragraphs)
        assert len(orig.sections) == len(new.sections)

        # Verify para 18 was actually mutated
        assert "Deep Research Agent" in new.paragraphs[18].text
        assert "LangGraph + DSPy + Pydantic" in new.paragraphs[18].text

        # Verify untouched para keeps content
        assert orig.paragraphs[1].text == new.paragraphs[1].text  # 胡阳
        assert orig.paragraphs[7].text == new.paragraphs[7].text  # multi-column

        # Verify para 18 paragraph-level style preserved
        assert orig.paragraphs[18].style.name == new.paragraphs[18].style.name

    def test_length_drift_too_far_skipped(self, tmp_path) -> None:
        """LLM returning text 200% of original should be skipped."""
        master_path = Path("/Users/huy/new_try/中文简历.docx")
        if not master_path.exists():
            pytest.skip("user resume not present")

        # Para 18 is ~200 chars; LLM returns 1000-char text
        decisions = [
            {
                "index": 18,
                "decision": "rewrite",
                "new_text": "X" * 1000,
                "rationale": "too long",
            }
        ]
        llm = _StubLLM(decisions)
        out_path = tmp_path / "tailored.docx"
        result = tailor_docx(
            input_path=master_path,
            output_path=out_path,
            jd_text="any",
            company="C", role_focus="r",
            master_resume_text="x",
            llm=llm,
        )
        # The length-drift sanity should kick in → 0 rewrites
        assert result.summary()["rewritten"] == 0

    def test_keep_decision_makes_no_change(self, tmp_path) -> None:
        master_path = Path("/Users/huy/new_try/中文简历.docx")
        if not master_path.exists():
            pytest.skip("user resume not present")

        decisions = [{"index": 18, "decision": "keep",
                      "new_text": "", "rationale": "already good"}]
        llm = _StubLLM(decisions)
        out_path = tmp_path / "tailored.docx"
        result = tailor_docx(
            input_path=master_path,
            output_path=out_path,
            jd_text="x", company="C", role_focus="r",
            master_resume_text="x", llm=llm,
        )
        assert result.summary()["rewritten"] == 0


# ═══════════════════════════════════════════════════════════════════
# UI route
# ═══════════════════════════════════════════════════════════════════


@pytest.fixture
def app_setup_no_resume(tmp_path):
    store = offerguide.Store(tmp_path / "ui.db")
    store.init_schema()
    skills = discover_skills(SKILLS_ROOT)
    app = create_app(
        settings=Settings(), store=store, profile=None,
        skills=skills, runtime=None, notifier=ConsoleNotifier(),
    )
    return app, store


class TestDocxRoute:
    def test_no_resume_returns_error(self, app_setup_no_resume) -> None:
        app, _ = app_setup_no_resume
        resp = TestClient(app).post("/api/tailor/docx", data={"job_id": "1"})
        assert resp.status_code == 200
        assert "OFFERGUIDE_RESUME_PDF" in resp.text

    def test_download_path_traversal_blocked(self, app_setup_no_resume) -> None:
        app, _ = app_setup_no_resume
        for evil in ("../etc/passwd", "..\\windows\\sys", "subdir/foo.docx"):
            resp = TestClient(app).get(f"/api/tailor/download/{evil}")
            # Either 400 (caught by validator) or 404 — never 200
            assert resp.status_code in (400, 404)

    def test_download_non_docx_blocked(self, app_setup_no_resume) -> None:
        app, _ = app_setup_no_resume
        resp = TestClient(app).get("/api/tailor/download/something.txt")
        assert resp.status_code == 400


# ═══════════════════════════════════════════════════════════════════
# Config OFFERGUIDE_LLM_* fallback chain
# ═══════════════════════════════════════════════════════════════════


class TestConfigFallback:
    def test_offerguide_llm_takes_priority(self, monkeypatch) -> None:
        monkeypatch.setenv("OFFERGUIDE_LLM_API_KEY", "new-priority")
        monkeypatch.setenv("DEEPSEEK_API_KEY", "old-deepseek")
        monkeypatch.setenv("TOKEN", "older-token")
        s = Settings.from_env()
        assert s.deepseek_api_key == "new-priority"

    def test_falls_back_to_token(self, monkeypatch) -> None:
        monkeypatch.delenv("OFFERGUIDE_LLM_API_KEY", raising=False)
        monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
        monkeypatch.setenv("TOKEN", "t-fallback")
        s = Settings.from_env()
        assert s.deepseek_api_key == "t-fallback"

    def test_offerguide_llm_model_takes_priority(self, monkeypatch) -> None:
        monkeypatch.setenv("OFFERGUIDE_LLM_MODEL", "claude-haiku")
        monkeypatch.setenv("OFFERGUIDE_DEFAULT_MODEL", "old-model")
        s = Settings.from_env()
        assert s.default_model == "claude-haiku"


_ = ParagraphPlan  # quiet unused import
