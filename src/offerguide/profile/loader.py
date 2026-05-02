"""Resume file → UserProfile.

Supports .pdf (via pypdf) and .docx (via python-docx). Auto-dispatches by
extension. The same env var ``OFFERGUIDE_RESUME_PDF`` accepts both — the
name is historical (W1 was PDF-only); .docx support added in W12-fix-c
when the user gave us a real ``中文简历.docx``.

W1 only does text extraction. The LLM-driven structured parse (filling out
`education`, `experience`, `skills`, `preferences`) lands in W2 once the LLM
client and the `parse_resume` SKILL are wired in.
"""

from __future__ import annotations

from pathlib import Path

from .schema import UserProfile


def load_resume_pdf(resume_path: str | Path) -> UserProfile:
    """Load a resume file (PDF or DOCX) and return UserProfile.

    Dispatches by file extension:
    - ``.pdf`` → pypdf text extraction
    - ``.docx`` → python-docx paragraph + table text extraction

    Whitespace is normalized at the boundaries only — inner formatting is
    preserved so the W2 LLM parser sees the original layout.
    """
    resume_path = Path(resume_path)
    if not resume_path.exists():
        raise FileNotFoundError(f"Resume file not found: {resume_path}")

    suffix = resume_path.suffix.lower()
    if suffix == ".pdf":
        raw_text = _extract_pdf_text(resume_path)
    elif suffix in (".docx", ".doc"):
        raw_text = _extract_docx_text(resume_path)
    else:
        raise ValueError(
            f"Unsupported resume format: {suffix} (need .pdf or .docx)"
        )

    return UserProfile(
        raw_resume_text=raw_text,
        source_pdf=str(resume_path.resolve()),
    )


def _extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip()).strip()


def _extract_docx_text(path: Path) -> str:
    """Pull text from .docx paragraphs + table cells.

    Preserves paragraph order; normalizes runs of whitespace within each
    paragraph (Word splits styling boundaries into runs which can leave
    weird spacing artifacts in the joined string).
    """
    import re as _re

    from docx import Document
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            # Collapse internal multi-space (Word run boundaries) to single
            text = _re.sub(r"[ \t]{2,}", "  ", text)
            parts.append(text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    parts.append(ct)
    return "\n".join(parts).strip()
