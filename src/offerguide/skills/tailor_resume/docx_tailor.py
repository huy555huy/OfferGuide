"""DOCX-level resume tailoring — preserves Word formatting end-to-end.

Why this is a separate module from the SKILL:
The ``tailor_resume`` SKILL outputs a markdown document with a change_log.
That works for *displaying* edits but not for **handing the user a real
.docx file they can submit**. Word resumes carry inline formatting (font,
size, bold, color, paragraph alignment, list bullets, page break, header)
that markdown can't roundtrip.

This module operates directly on python-docx ``Document`` objects:

1. **Read** original .docx → build a paragraph-level edit plan (which
   paragraphs are eligible for rewrite, which are skipped)
2. **Send to LLM** the eligible paragraphs + JD + master resume, get back
   per-paragraph rewrites (or "keep as-is")
3. **Write back** by mutating ``run.text`` in place — paragraph-level
   styles (font, alignment, bullet) are preserved automatically because
   we don't touch the run's style attributes

Skip rules (what NOT to rewrite — protect format integrity):
- Empty paragraphs
- Header / name lines (paragraphs < 8 chars + bold)
- Multi-column layout paragraphs (contain ``\\t`` or 3+ consecutive
  spaces — typical of "学校 ... 时间" two-column rows)
- Date / phone number / email lines (high regex match)

The "runs collapse" question:
A paragraph with 14 runs (Word's natural fragmentation by inline style
boundaries) will, after rewrite, end up with the new text in run[0] and
runs[1:] cleared to "". This means **inline styling within the paragraph
collapses to whatever style run[0] has**. For Chinese resumes this is
usually fine — paragraph-level styling (`paragraph.style`, alignment,
indent) is what carries the visual weight; inline italic/color shifts
inside a paragraph are rare in 简历.

Tested on user's actual 中文简历.docx (W12-fix-c, 2026-05-02).
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from ...llm import LLMClient, LLMError

log = logging.getLogger(__name__)


@dataclass
class ParagraphPlan:
    """Per-paragraph eligibility decision + LLM result."""
    index: int
    """Paragraph index in the doc (0-based)."""

    original_text: str
    eligible: bool
    skip_reason: str = ""
    """If not eligible, why — surfaced in the change_log so user sees we
    *intentionally* didn't touch sensitive areas (name, dates, etc)."""

    new_text: str = ""
    """LLM-rewritten text. Empty → no change applied."""

    rationale: str = ""
    """Why this rewrite (referencing JD line / ATS keyword)."""


@dataclass
class DocxTailorResult:
    """Outcome of docx_tailor — used by the route to render summary + serve file."""
    output_path: Path
    paragraph_plans: list[ParagraphPlan] = field(default_factory=list)

    @property
    def changes(self) -> list[ParagraphPlan]:
        """Plans that actually mutated the doc (eligible + new_text differs)."""
        return [
            p for p in self.paragraph_plans
            if p.eligible and p.new_text and p.new_text != p.original_text
        ]

    @property
    def skipped(self) -> list[ParagraphPlan]:
        """Plans that were intentionally not rewritten."""
        return [p for p in self.paragraph_plans if not p.eligible]

    def summary(self) -> dict[str, int]:
        return {
            "total_paragraphs": len(self.paragraph_plans),
            "eligible": sum(1 for p in self.paragraph_plans if p.eligible),
            "rewritten": len(self.changes),
            "skipped": len(self.skipped),
        }


# ─────────── eligibility rules ─────────────────────────────


_EMAIL_PHONE_DATE_RE = re.compile(
    r"\b\d{2,4}[-/年]\d{1,2}[-/月]?(\d{1,2})?日?\b"  # date
    r"|\b\d{4,11}\b"                                   # phone
    r"|@[\w.]+\.\w{2,5}",                              # email
)
# Multi-column = tab OR 5+ consecutive spaces. (3+ would catch normal
# Word run-boundary artifacts like "...预设阶段的   workflow agent...".)
_MULTI_COL_RE = re.compile(r"\t| {5,}")


def _is_eligible(paragraph_text: str, run_count: int, is_bold: bool) -> tuple[bool, str]:
    """Decide whether to send this paragraph to the LLM for rewrite.

    Returns (eligible, skip_reason). Skip is intentional, not a failure —
    we protect name / dates / multi-column layout from being touched.
    """
    text = (paragraph_text or "").strip()
    if not text:
        return False, "empty"

    # Header / name line — short bold text
    if is_bold and len(text) < 8:
        return False, "section_header"

    # Multi-column layout (tab / 3+ spaces — typical of 学校 + 时间 rows)
    if _MULTI_COL_RE.search(paragraph_text):
        return False, "multi_column_layout"

    # Date / phone / email line — must not paraphrase
    if _EMAIL_PHONE_DATE_RE.search(text):
        return False, "contains_date_phone_email"

    # Very short non-bold (likely sub-heading or label like "教育背景")
    if len(text) < 12:
        return False, "too_short"

    return True, ""


# ─────────── LLM rewrite step ─────────────────────────────


_REWRITE_PROMPT = """你是 OfferGuide 的简历段落微调员。给定:
- 用户 master 简历全文 (上下文用)
- 目标 JD 全文
- 一组待改写的段落 (按 index)

返回 JSON 数组, 每个元素对应一个段落:
{{
  "index": <int, 必填, 与输入对齐>,
  "decision": "rewrite" | "keep" ,
  "new_text": <str, decision=rewrite 时必填; keep 时填空字符串>,
  "rationale": <str, 一句话, 引用 JD 第 X 条 / ATS 关键词 / profile 必备项>
}}

**严禁** (ANY 违规这条段落必须 decision=keep):
1. 新增 master_resume 没有的实习公司/项目/比赛/论文
2. 修改可被验证的硬事实 (学校 / 学位 / 实习时间 / 数字结果)
3. 夸大或缩小数字 (AUC 0.04 不能变 5%)
4. 编造 master_resume 没提到的技术栈
5. 修改段落字符长度 ±50% 之外 (会破坏 Word 排版)

**鼓励** (有这些信号就 rewrite):
- master_resume 写得太啰嗦 / 不贴 JD 关键词 → 改措辞贴 JD
- 同一个项目可以强调不同侧面以匹配 JD → 调整侧重
- 漏写但简历其他地方有证据的 ATS 关键词可以加进段落

**重要**: new_text 字符数和 original 不能差太远 (Word 段落长度变化会跳行重排)。
理想: ±20% 之内。

只返回 JSON 数组。不要 markdown 代码块。
"""


def _rewrite_with_llm(
    *,
    eligible_plans: list[ParagraphPlan],
    master_resume: str,
    jd_text: str,
    company: str,
    role_focus: str,
    llm: LLMClient,
) -> list[dict[str, Any]]:
    """Send eligible paragraphs to LLM, return per-paragraph decisions."""
    if not eligible_plans:
        return []

    paragraphs_block = "\n\n".join(
        f"[index={p.index}] {p.original_text}"
        for p in eligible_plans
    )
    user_msg = (
        f"【目标公司】{company}\n"
        f"【角色聚焦】{role_focus}\n\n"
        f"【JD 全文】\n{jd_text[:3000]}\n\n"
        f"【master 简历全文 (改写时不要超出此处提到的事实)】\n"
        f"{master_resume[:3000]}\n\n"
        f"【待决策的段落】\n{paragraphs_block}"
    )

    try:
        resp = llm.chat(
            messages=[
                {"role": "system", "content": _REWRITE_PROMPT},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.0,
            json_mode=True,
        )
    except LLMError as e:
        log.warning("docx_tailor LLM failed: %s", e)
        return []

    try:
        decisions = json.loads(resp.content)
    except json.JSONDecodeError:
        log.warning("docx_tailor LLM returned non-JSON")
        return []

    if not isinstance(decisions, list):
        return []
    return decisions


# ─────────── main entry point ─────────────────────────────


def tailor_docx(
    *,
    input_path: str | Path,
    output_path: str | Path,
    jd_text: str,
    company: str,
    role_focus: str,
    master_resume_text: str,
    llm: LLMClient,
) -> DocxTailorResult:
    """Read a .docx resume, rewrite eligible paragraphs, write a new .docx.

    Format preservation strategy:
    - Paragraph styles, alignment, indent, bullets — all preserved
      (we never touch ``paragraph.style`` / paragraph properties)
    - Within each paragraph we collapse all run.text into runs[0].text
      and clear runs[1:].text — inline style boundaries within a paragraph
      are lost but paragraph-level styling is intact
    - Skipped paragraphs (header / dates / multi-column rows) are
      untouched at the run level
    """
    from docx import Document

    input_path = Path(input_path)
    output_path = Path(output_path)
    if not input_path.exists():
        raise FileNotFoundError(f"resume not found: {input_path}")

    doc = Document(str(input_path))

    # Step 1: build per-paragraph plans
    plans: list[ParagraphPlan] = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text or ""
        run_count = len(p.runs)
        is_bold = any(r.bold for r in p.runs if r.bold is not None)
        eligible, skip_reason = _is_eligible(text, run_count, is_bold)
        plans.append(ParagraphPlan(
            index=i,
            original_text=text,
            eligible=eligible,
            skip_reason=skip_reason,
        ))

    # Step 2: send eligible paragraphs to LLM
    eligible = [p for p in plans if p.eligible]
    decisions = _rewrite_with_llm(
        eligible_plans=eligible,
        master_resume=master_resume_text,
        jd_text=jd_text,
        company=company,
        role_focus=role_focus,
        llm=llm,
    )

    # Step 3: apply decisions to plans
    decisions_by_idx: dict[int, dict[str, Any]] = {}
    for d in decisions:
        if isinstance(d, dict) and isinstance(d.get("index"), int):
            decisions_by_idx[d["index"]] = d
    for plan in plans:
        d = decisions_by_idx.get(plan.index)
        if not d:
            continue
        if d.get("decision") != "rewrite":
            continue
        new_text = (d.get("new_text") or "").strip()
        if not new_text:
            continue
        # Sanity: new text length within ±50% of original (compare stripped
        # to stripped — Word paragraphs may have trailing whitespace)
        orig_stripped = len(plan.original_text.strip())
        new_stripped = len(new_text.strip())
        if orig_stripped > 0 and not (0.4 <= new_stripped / orig_stripped <= 1.6):
            log.info(
                "docx_tailor: skip rewrite of paragraph %d — length drift %d→%d",
                plan.index, orig_stripped, new_stripped,
            )
            continue
        plan.new_text = new_text
        plan.rationale = (d.get("rationale") or "")[:200]

    # Step 4: write back to docx — runs collapse to runs[0]
    for plan in plans:
        if not plan.eligible or not plan.new_text:
            continue
        if plan.new_text == plan.original_text:
            continue
        para = doc.paragraphs[plan.index]
        if not para.runs:
            # No runs to collapse — just add text via add_run (rare for non-empty para)
            para.add_run(plan.new_text)
            continue
        # Put new text in run[0], clear the rest
        para.runs[0].text = plan.new_text
        for r in para.runs[1:]:
            r.text = ""

    # Step 5: save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return DocxTailorResult(output_path=output_path, paragraph_plans=plans)
